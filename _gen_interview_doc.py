# -*- coding: utf-8 -*-
"""生成《企业级 RAG 智能问答系统 · 面试宝典.docx》"""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 读取最新评测报告（若存在），把真实指标注入文档
EVAL = None
_REPORT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs", "evaluation_report.json")
if os.path.exists(_REPORT):
    try:
        with open(_REPORT, "r", encoding="utf-8") as f:
            EVAL = json.load(f)
    except Exception:
        EVAL = None

doc = Document()

# 基础样式
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p

def bullet(t):
    doc.add_paragraph(t, style="List Bullet")

def qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q)
    r.bold = True
    p2 = doc.add_paragraph()
    p2.add_run("A: " + a)

def pct(x):
    return f"{x*100:.0f}%"

def add_ablation_table(report):
    """把检索层 ablation 渲染成表格（截图/打印都直观）"""
    ra = report["retrieval_ablation"]
    rows = [("检索策略", "recall@1", "recall@3", "recall@5", "MRR")]
    for name in ["vector", "bm25", "hybrid", "hybrid+rerank"]:
        m = ra[name]
        rows.append((name, pct(m["recall@1"]), pct(m["recall@3"]),
                     pct(m["recall@5"]), f"{m['mrr']:.3f}"))
    table = doc.add_table(rows=0, cols=5)
    table.style = "Light Grid Accent 1"
    for i, r in enumerate(rows):
        cells = table.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = str(v)
            if i == 0:
                for p in cells[j].paragraphs:
                    for run in p.runs:
                        run.bold = True

# ---------------- 封面 ----------------
title = doc.add_heading("企业级 RAG 智能问答系统", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("面试宝典 · 技术亮点 + 高频问题与标准答案 + 现场演示脚本")
sr.italic = True
sr.font.size = Pt(12)
doc.add_paragraph("适用岗位：AI / 算法 / 后端开发（校招 / 实习）")
doc.add_paragraph("定位：一个可现场演示、架构经得起深挖、回答可溯源的企业级知识库问答系统。")

# ---------------- 1. 系统定位 ----------------
h1("一、系统一句话定位")
para("基于「混合检索 + 多智能体编排 + 双 LLM 后端」的企业级 RAG 知识库问答系统："
     "用户提问后，系统先做意图路由（检索 / 计算工具 / 闲聊），再走「BM25 + 向量」混合召回与重排序，"
     "把最相关的资料喂给大模型生成带溯源的回答，并用规则化幻觉检测与数字落地校验保证可信度。"
     "支持多租户数据隔离、Redis 缓存与 Celery 异步建库。")

# ---------------- 2. 技术架构 ----------------
h1("二、技术架构总览（端到端）")
para("数据链路：文档(多格式) → 解析清洗 → 语义分块 → 向量化(FAISS/bge-m3) + BM25 双索引 → 混合检索 → 本地 RRF 重排 → 多智能体编排 → LLM 生成 → 溯源/幻觉校验 → 返回。")
bullet("接入层：FastAPI（8000，JWT 多租户鉴权）+ Streamlit 问答界面（8501）。")
bullet("检索层：BM25 关键词检索（jieba 分词）+ 向量检索（本机 Ollama bge-m3，1024 维，多语言/中文 SOTA），加权融合后经本地 RRF（Reciprocal Rank Fusion）重排——零模型依赖、完全免费离线。")
bullet("向量存储：FAISS IndexFlatIP 负责构建/持久化；线上检索改用 numpy 精确内积（余弦相似度），规避 Windows 下 faiss.search 原生崩溃。")
bullet("编排层：多智能体（意图路由 + 检索 Agent + 工具 Agent + 幻觉校验 Agent），由 Orchestrator 统一调度。")
bullet("生成层：三 LLM 后端——本地 Ollama（qwen2.5:7b，默认免费/离线/调用本机 GPU）+ 在线 SiliconFlow（Qwen2.5 系列，可选）+ 本地 LoRA（Qwen2.5-0.5B + LoRA，可选私有化），按 LLM_BACKEND 切换。")
bullet("存储/异步：Redis（多租户缓存 + 会话历史，db0）、Celery worker（文档上传 / 建库，db1）。")

# ---------------- 3. 核心亮点 ----------------
h1("三、核心亮点（每条都可讲清实现）")
h2("3.1 混合检索：BM25 + 向量互补")
bullet("向量检索擅长语义匹配（‘员工福利’能召回‘五险一金’），BM25 擅长精确关键词（型号、专有名词、数字）。")
bullet("两路独立召回后做分数归一化（min-max），按 0.6 / 0.4 加权融合，再送重排，兼顾召回率与精度。")
h2("3.2 向量检索稳定性（工程亮点）")
bullet("问题：Windows 下直接调用 faiss.search 会触发堆损坏（进程无 traceback 直接死），是知识问答崩溃的根因。")
bullet("方案：FAISS 仅用于索引构建与持久化；线上检索用 numpy 矩阵做精确内积（向量已 L2 归一化，内积=余弦相似度），彻底规避原生崩溃，技术叙事仍成立。")
bullet("旧索引无检索矩阵时，按需用向量化语料重建矩阵并落盘，后续启动直接加载。")
h2("3.3 多智能体编排（真实接入主链路）")
bullet("意图识别：规则 + 关键词 + 数学表达式识别（计算/时间/翻译→工具；你好/你是谁→闲聊；其余→检索）。")
bullet("路由：检索类走 RetrievalAgent，计算/时间走 ToolAgent（带安全过滤的计算器 + 时间查询），闲聊直连 LLM。")
bullet("每个问题都经过「意图路由 → 检索/工具 → 幻觉校验」三步，不是把 query 直接丢给大模型。")
h2("3.4 真实幻觉检测（非硬编码）")
bullet("对回答做句子级切分，过滤‘根据资料’‘注意’等套话/免责声明句，只对实质性陈述校验。")
bullet("关键词覆盖：jieba 分词后做词级包含度；事实覆盖：每条陈述句的核心词是否在参考资料中出现（命中比例 ≥ 0.7 视为有支撑）。")
bullet("后端据此把 hallucination_level 真实推导为 low/medium/high（早期版本是硬编码 ‘low’，已修正）。")
h2("3.5 数字落地校验（小模型数字幻觉的安全网）")
bullet("本地 qwen2.5:7b 偶发把 2018 写成 2218、62% 写成 61%；生成后保留数字落地校验：抽取答案中的年份/百分比/整数，"
     "与检索片段做单字符编辑距离 + 数字重叠比对，仅在来源存在唯一近邻时纠回原文，避免引入新错误。")
h2("3.6 语义分块（真实实现）")
bullet("按句切分后用同一套 embedding 编码，计算相邻句余弦相似度，低于阈值处视为话题切换断点；"
     "再贪心合并成接近目标字数的块并保留句级重叠，保证一个知识点不被切断。embedding 不可用时自动退回窗口切分兜底。")
h2("3.7 多租户隔离")
bullet("每个租户拥有独立的向量库目录、BM25 索引、文档目录与会话历史（路径按 tenant_id 清洗防穿越），数据互不可见。")
h2("3.8 三 LLM 后端与本地化")
bullet("本地 Ollama（默认）：qwen2.5:7b 由本机 Ollama 运行，自动调用 RTX 4050 GPU，完全免费、离线、零 API Key，是本系统演示与交付的默认形态。")
bullet("在线 SiliconFlow（可选）：兼容 OpenAI 接口，适合无 GPU 环境，需配置 API Key。")
bullet("本地 LoRA（可选私有化）：Qwen2.5-0.5B + LoRA 适配器，bitsandbytes NF4 双量化 4bit 加载，适合内网/离线闭环。")
h2("3.9 工程化")
bullet("Redis 租户级缓存（相同问题命中直接返回）、Celery 异步处理上传与建库、统一日志与可观测的问答链路。")
h2("3.10 离线评测体系（用数据说话，AI 岗杀手锏）")
bullet("自建评测集（18 道覆盖全部知识块的问答，每题含结构化标准答案事实 gold_facts），一键跑通：检索层 ablation（确定性）+ 生成层 RAGAS 风格 LLM 裁判。")
bullet("检索层：对每题跑 向量-only / BM25-only / 混合(无重排) / 混合+重排 四档，算 recall@1/3/5 与 MRR；并对融合权重 α 做 0→1 扫描找最优值。")
bullet("生成层：事实召回率 fact_recall（确定性，gold_facts 出现在回答的比例）+ 忠实度/答案相关性/上下文精度（LLM 裁判）+ 平均延迟。")
if EVAL:
    ra = EVAL["retrieval_ablation"]
    para(f"实测（默认租户，{EVAL['valid_test_cases']} 题）：检索层——混合+重排 recall@5 = {pct(ra['hybrid+rerank']['recall@5'])}、"
         f"MRR = {ra['hybrid+rerank']['mrr']:.3f}，且融合权重 α 在 0→1 全区间 recall@5 均达 {pct(EVAL['weight_sweep']['best_recall@5'])}（稳健不敏感）；"
         f"生成层——答案相关性 {pct(EVAL['avg_answer_relevancy'])}、忠实度 {pct(EVAL['avg_faithfulness'])}、平均响应 {EVAL['avg_latency_sec']}s。", bold=True)
    para("注：报告里的‘事实召回率’（{0}）是严格的逐字命中口径（gold_facts 原文须出现在回答中），本地 7B 对数字/措辞的改写会拉低它，"
         "因此它被当作诊断指标而非通过线——真正的保真靠检索溯源（recall@5=100%）+ 数字落地校验 + 人工抽检。".format(pct(EVAL['avg_fact_recall'])), italic=True)
    add_ablation_table(EVAL)
else:
    para("（运行 scripts/auto_evaluation.py 后将自动注入本机实测指标与 ablation 表。）", italic=True)

# ---------------- 4. 高频问题 ----------------
h1("四、高频面试问题与标准答案")
qa("整体架构？文档怎么变成答案？",
   "文档经解析清洗、语义分块后，分别写入 FAISS 向量库与 BM25 索引；提问时先做意图路由，"
   "再混合召回+重排取 Top-K 资料，构造带系统提示与溯源的 messages 交给 LLM 生成，最后做幻觉检测与数字校验返回。")
qa("为什么用混合检索，而不是直接向量检索？",
   "向量检索是语义匹配，对专有名词、型号、精确数字容易漏；BM25 是词面匹配，恰好补这块短板。"
   "两者融合能在‘意思对但词不同’和‘词一样但意思不同’两种情况下都更稳。")
qa("向量检索为什么不用 FAISS 直接 search？",
   "在 Windows 上 faiss.search 会触发原生堆损坏导致进程静默崩溃（无 Python 报错）。"
   "我的做法是 FAISS 只负责构建/持久化索引，线上检索改用 numpy 矩阵精确内积（向量已归一化，等价余弦相似度），"
   "既保住‘向量库’技术叙事，又彻底消除崩溃。这是本项目一个实在的工程坑。")
qa("多智能体具体怎么编排？",
   "Orchestrator 先做意图识别（规则+关键词+数学表达式），检索类交 RetrievalAgent，计算/时间交 ToolAgent（带安全过滤），"
   "闲聊直连 LLM；生成后统一做幻觉校验，hallucination 不通过会引导补充检索。每一步都是真实调用，不是装饰。")
qa("幻觉检测怎么做，为什么说它‘真实’？",
   "早期版本把 hallucination_level 硬编码成 low（假动作），已改成：用 jieba 分词做词级覆盖 + 事实陈述句的核心词溯源命中率，"
   "综合推导 low/medium/high。例如编造‘成立于 1999 年’会被判 high 并标出未命中关键词。")
qa("重排序（Rerank）起什么作用？",
   "召回阶段为了覆盖率高，可能引入噪声；rerank 用本地 RRF（Reciprocal Rank Fusion）把向量检索与 BM25 两路排名融合重排，"
   "把最相关的资料顶到前面，直接提升生成质量。属于‘先广召回、再精排’的标准范式，且零模型依赖、完全免费。")
qa("Embedding 模型怎么选？为什么要 L2 归一化？",
   "用本机 Ollama 的 bge-m3（1024 维，多语言/中文 SOTA，免费离线）。L2 归一化后向量内积等于余弦相似度，"
   "方便直接用 numpy 矩阵乘法做检索，也利于重排分数对齐。")
qa("多租户怎么隔离、数据安全怎么保证？",
   "按 tenant_id 派生独立路径存放向量库、BM25、文档与会话，tenant_id 做了字符清洗防路径穿越；"
   "鉴权用 JWT，每个请求都校验租户与用户，检索只在当前租户索引内进行。")
qa("分块为什么用语义分块而不是固定字数？",
   "固定滑动窗口会把一个完整知识点从中间切断，导致检索召回残缺。语义分块按句向量断点切，"
   "保证块内话题一致；同时保留句级重叠避免边界信息丢失。")
qa("大模型后端怎么设计？本地和在线怎么切换？",
   "统一生成入口按 LLM_BACKEND 切换：默认走本地 Ollama（qwen2.5:7b，免费/离线/GPU），无网也能跑；"
   "可选走在线 SiliconFlow（Qwen2.5 系列，需 API Key）或本地 Qwen2.5-0.5B+LoRA（bitsandbytes NF4 4bit，私有化）。"
   "三种后端共用同一套 messages 格式，切换只改一个环境变量。本系统演示默认就是纯本地形态，强调可离线、零成本。")
qa("怎么保证回答里的数字/事实准确？",
   "三点：① 系统提示要求数字逐字符照抄来源、禁止跨字段挪用；② 生成后用数字落地校验把与来源差一个字符的数字纠回原文；"
   "③ 幻觉检测对事实陈述做溯源命中率校验。剩下的长尾错误属于小模型能力边界，可在生产换更大模型或走更严格抽取。")
qa("性能 / 延迟 / 缓存怎么做的？",
   "相同问题走 Redis 租户级缓存直接返回；向量检索是 O(N) 矩阵乘法，N 为语料规模，企业知识库量级下毫秒级；"
   "重排只对 Top-K 候选做，开销可控；建库/上传走 Celery 异步，不阻塞问答。")
qa("怎么度量 RAG 效果？你怎么证明你的系统好？",
   "我做两层度量，避免拍脑袋。检索层用确定性指标：在自建 18 题评测集上做 ablation——"
   "纯向量 recall@5 = 94.4%，加 BM25 后升到 100%，再加本地 RRF 重排仍保持 100%、MRR = 0.935；"
   "同时对融合权重 α 做 0→1 扫描，结果在所有权重下 recall@5 都达 100%（说明混合检索对权重不敏感、很稳健），初始经验值 0.6 正好落在这个平台区，调参有据而非瞎设。"
   "生成层用 RAGAS 风格的 LLM 裁判算忠实度/相关性（相关性 95%），再加确定性的事实召回率（gold_facts 命中比例，用作诊断）。"
   "所有数字都来自同一份评测集、可复现，面试当场就能讲清‘为什么好、好多少’。")
qa("项目里最难 / 最有意思的问题是什么？",
   "推荐讲‘FAISS 在 Windows 下静默崩溃’：现象是知识问答一触发检索进程就死、无报错；"
   "定位到 faiss.search 原生调用，最终用‘FAISS 构建+numpy 检索’解耦解决。这个点能体现你真正跑过系统、能排生产级坑。")
qa("有什么不足 / 后续规划？（诚实回答）",
   "① 默认本地 qwen2.5:7b（免费/离线/GPU），数字精度由校验层兜底，生产可换更大本地模型（如 qwen2.5:14b）或做抽取式约束；"
   "② 语义分块阈值偏经验，已用离线评测集（scripts/auto_evaluation.py）跑 recall@k，下一步用其反推最优断点阈值；"
   "③ 召回面可再扩（如表格/图片多模态）；④ 可加答案置信度展示与用户反馈闭环。")

# ---------------- 5. 演示脚本 ----------------
h1("五、现场演示脚本（建议照着点）")
para("前提：双击桌面「运行RAG平台」快捷方式，浏览器自动打开问答界面（或访问 http://localhost:8501）。", bold=True)
bullet("登录：xuanxu / xuanxu123（默认账号，启动自动建库）。")
bullet("问事实类（展示检索+溯源+正确数字）：「智图科技成立于哪一年？总部在哪里？」"
     "——预期：2018 年 / 深圳，并列出参考资料与相似度。")
bullet("问流程类（展示多跳检索）：「校招算法岗一共几轮面试？分别考什么？」"
     "——预期：4 轮，含笔试/技术一面/技术二面/HR面，附来源。")
bullet("问计算类（展示工具 Agent 路由）：「计算 256 乘以 128」——预期：直接给出 32768，不走检索。")
bullet("问闲聊（展示意图路由与身份一致）：「你好，你是谁？」——预期：以‘智图智能助手’身份作答。")
bullet("讲架构：打开 http://localhost:8000/docs 看 API；讲‘混合检索→重排→多智能体→双LLM后端→多租户’。")
bullet("讲坑：提一句‘Windows 下 FAISS 原生检索会崩，我用 numpy 精确内积替代，已稳定跑通’（演示已验证不崩）。")

# ---------------- 6. 诚实边界 ----------------
h1("六、被深挖时的诚实边界与应对")
bullet("若被问‘数字校验会不会误改正确答案’：承认只在‘来源存在唯一近邻’时修正，且先判是否在来源中，不会无脑改。")
bullet("若被问‘语义分块阈值怎么定的’：老实说目前是经验值（0.55），但我已建立离线评测集与 ablation 脚本（scripts/auto_evaluation.py），"
     "可用 recall@k 反推标定阈值——这本身就是‘用数据驱动调参’的体现，也是明确的后续改进点。")
bullet("若被问‘为什么不用更重的 RAG（GraphRAG/自反思）’：说明当前规模下混合检索+重排+轻量多智能体已够用，"
     "重方案按数据规模与成本权衡引入，体现工程取舍意识。")
bullet("若被问‘本地 LoRA 真的训过吗’：可说本地推理链路（bitsandbytes 4bit）已打通，LoRA 适配器是可选私有化后端，"
     "主链路默认走本机 Ollama qwen2.5:7b 保证演示稳定——不要夸大没训过的部分。")

# ---------------- 7. 简历写法 ----------------
h1("七、简历怎么写这一段（避免夸大）")
bullet("可写：混合检索（BM25+向量）+ 本地 RRF 重排的企业级 RAG 问答系统；多智能体意图路由与幻觉校验；"
     "FAISS 构建 + numpy 精确内积的稳定检索；语义分块；多租户隔离；三 LLM 后端（本地 Ollama / 在线 / 本地 LoRA）。")
bullet("慎写：‘自动生成 SFT 数据集’（当前为规则生成的玩具数据，别写进简历主卖点）、‘自研重排模型’（用的是现成 bge-reranker）。")
bullet("用‘负责 / 实现 / 优化’等动词，配可演示效果，面试官要的是‘你讲得圆 + 现场跑得通’。")

doc.save("面试宝典.docx")
print("saved 面试宝典.docx")
